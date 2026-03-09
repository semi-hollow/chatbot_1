"""生成第一阶段示例 Word 文档（避免将二进制文件直接提交到仓库）。"""
from pathlib import Path
from docx import Document


def main() -> None:
    target = Path("data/documents/settlement_checklist.docx")
    target.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("结算检查清单（示例）", level=1)
    doc.add_paragraph("1. 核对交易流水与台账。")
    doc.add_paragraph("2. 核对费用与税金。")
    doc.add_paragraph("3. 输出异常项并确认责任人。")
    doc.save(target)
    print(f"已生成: {target}")


if __name__ == "__main__":
    main()
